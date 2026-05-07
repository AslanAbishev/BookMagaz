"""Export the experimental engineering markdown report to a simple DOCX file."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = PROJECT_ROOT / "docs" / "ASSIGNMENT_3_EXPERIMENTAL_REPORT.md"
OUTPUT = PROJECT_ROOT / "docs" / "Assignment 3 Experimental Engineering Report.docx"


def add_table(document: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 2:
        return

    headers = rows[0]
    data_rows = [row for row in rows[2:] if len(row) == len(headers)]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in data_rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph("")


def build_docx() -> Path:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Assignment 3: Experimental Engineering")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("GoodBooks Performance, Mutation, and Chaos Testing Report")

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            i += 1
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            paragraph = document.add_paragraph()
            paragraph.style = "Intense Quote"
            paragraph.add_run("\n".join(code_lines))
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue
        if line.startswith("- "):
            while i < len(lines) and lines[i].startswith("- "):
                document.add_paragraph(lines[i][2:], style="List Bullet")
                i += 1
            continue

        document.add_paragraph(line)
        i += 1

    document.add_page_break()
    document.add_heading("Screenshots To Insert Before Submission", level=1)
    for item in [
        "Terminal screenshot of performance experiment execution",
        "Terminal screenshot of mutation experiment execution and mutation score",
        "Terminal screenshot of chaos experiment execution",
        "Screenshot of performance_results.csv or chart version",
        "Screenshot of mutation_results.csv",
        "Screenshot of chaos_results.csv",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_docx()
    print(path)
